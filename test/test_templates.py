"""
Test script to diagnose template access issues
Run this to verify your template files are accessible
"""

import os
import sys
from pathlib import Path

def test_template_access():
    """Test if templates can be accessed and opened"""
    
    print("=" * 70)
    print("TEMPLATE ACCESS TEST")
    print("=" * 70)
    
    # Test 1: Find project root
    print("\n[TEST 1] Finding project root...")
    current_dir = "W:/competition/project_root/data/input/template"
    print(f"Current directory: {current_dir}")
    
    # Try to find project root
    potential_roots = [
        current_dir,
        os.path.dirname(current_dir),
        os.path.dirname(os.path.dirname(current_dir))
    ]
    
    project_root = None
    for root in potential_roots:
        if os.path.exists(os.path.join(root, "core")) or \
           os.path.exists(os.path.join(root, "ui")):
            project_root = root
            print(f"✅ Project root found: {project_root}")
            break
    
    if not project_root:
        print("⚠️  Could not find project root")
        project_root = current_dir
    
    # Test 2: Check template directories
    print("\n[TEST 2] Checking template directories...")
    
    template_paths = [
        os.path.join(project_root, "data", "templates"),
        os.path.join(project_root, "data", "input", "template"),
        os.path.join(project_root, "templates"),
    ]
    
    existing_dirs = []
    for path in template_paths:
        exists = os.path.exists(path)
        print(f"  {'✅' if exists else '❌'} {path}")
        if exists:
            existing_dirs.append(path)
    
    if not existing_dirs:
        print("\n❌ No template directories found!")
        return False
    
    # Test 3: List files in template directories
    print("\n[TEST 3] Listing template files...")
    
    template_files = {}
    for dir_path in existing_dirs:
        print(f"\n📁 {dir_path}")
        try:
            files = os.listdir(dir_path)
            docx_files = [f for f in files if f.endswith('.docx')]
            
            if not docx_files:
                print("  ⚠️  No .docx files found")
            else:
                for file in docx_files:
                    full_path = os.path.join(dir_path, file)
                    size = os.path.getsize(full_path)
                    print(f"  ✅ {file} ({size:,} bytes)")
                    
                    # Categorize
                    if 'convocation' in file.lower() or 'surveillance' in file.lower():
                        template_files['convocation'] = full_path
                    elif 'affectation' in file.lower() or 'enseignant' in file.lower():
                        template_files['affectation'] = full_path
        except Exception as e:
            print(f"  ❌ Error reading directory: {e}")
    
    # Test 4: Try to open templates with python-docx
    print("\n[TEST 4] Testing template access with python-docx...")
    
    try:
        from docx import Document
        print("✅ python-docx is installed")
    except ImportError:
        print("❌ python-docx not installed!")
        print("   Install with: pip install python-docx")
        return False
    
    # Test each template
    for template_type, template_path in template_files.items():
        print(f"\n  Testing {template_type}: {os.path.basename(template_path)}")
        
        try:
            # Try to open
            doc = Document(template_path)
            print(f"    ✅ File opens successfully")
            
            # Check structure
            print(f"    📊 Paragraphs: {len(doc.paragraphs)}")
            print(f"    📊 Tables: {len(doc.tables)}")
            
            # Check for expected content
            if template_type == 'convocation':
                found_placeholder = any("Mr/Mme" in p.text for p in doc.paragraphs)
                print(f"    {'✅' if found_placeholder else '⚠️ '} Contains 'Mr/Mme': {found_placeholder}")
                
                if doc.tables:
                    table = doc.tables[-1]
                    print(f"    📊 Last table: {len(table.rows)} rows × {len(table.columns)} cols")
                    if len(table.columns) >= 3:
                        print(f"    ✅ Table has 3+ columns (Date, Heure, Durée)")
                    else:
                        print(f"    ⚠️  Table has only {len(table.columns)} columns")
            
            elif template_type == 'affectation':
                found_au = any("AU :" in p.text or "AU:" in p.text for p in doc.paragraphs)
                found_date = any("Date :" in p.text or "Date:" in p.text for p in doc.paragraphs)
                print(f"    {'✅' if found_au else '⚠️ '} Contains 'AU :': {found_au}")
                print(f"    {'✅' if found_date else '⚠️ '} Contains 'Date :': {found_date}")
                
                if doc.tables:
                    table = doc.tables[-1]
                    print(f"    📊 Last table: {len(table.rows)} rows × {len(table.columns)} cols")
                    if len(table.columns) >= 3:
                        print(f"    ✅ Table has 3+ columns (Enseignant, Salle, Signature)")
                    else:
                        print(f"    ⚠️  Table has only {len(table.columns)} columns")
            
        except Exception as e:
            print(f"    ❌ Error opening file: {e}")
            import traceback
            print(traceback.format_exc())
    
    # Test 5: Try creating export function
    print("\n[TEST 5] Testing DocumentExporter initialization...")
    
    if not template_files:
        print("❌ No templates found to test with")
        return False
    
    template_dir = os.path.dirname(list(template_files.values())[0])
    
    try:
        # Import the module
        sys.path.insert(0, project_root)
        from core.export_docs import DocumentExporter
        
        print(f"✅ Successfully imported DocumentExporter")
        print(f"   Template dir: {template_dir}")
        
        # Try to initialize
        exporter = DocumentExporter(template_dir=template_dir)
        print(f"✅ DocumentExporter initialized successfully")
        print(f"   Convocation template: {os.path.basename(exporter.convocation_template)}")
        print(f"   Affectation template: {os.path.basename(exporter.affectation_template)}")
        
    except Exception as e:
        print(f"❌ Error initializing DocumentExporter: {e}")
        import traceback
        print(traceback.format_exc())
        return False
    
    # Summary
    print("\n" + "=" * 70)
    print("SUMMARY")
    print("=" * 70)
    
    if template_files:
        print("✅ Templates found and accessible")
        print(f"\n💡 Use this in your code:")
        print(f'\n   template_dir = r"{template_dir}"')
        print(f'\n   from core.export_docs import export_surveillance_documents')
        print(f'\n   results = export_surveillance_documents(')
        print(f'       df_assignments=df,')
        print(f'       output_dir="output",')
        print(f'       template_dir=template_dir')
        print(f'   )')
        return True
    else:
        print("❌ No usable templates found")
        return False


if __name__ == "__main__":
    success = test_template_access()
    sys.exit(0 if success else 1)