"""
Module for exporting surveillance assignments to Word documents.
Generates individual convocations and collective affectation lists.
"""

from docx import Document
from PyQt5.QtWidgets import QMessageBox
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
import sys
from typing import Dict, List
import pandas as pd
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from workers.db_handler import DBHandler
from helper.resource_path import resource_path

def get_project_root():
    """
    Get the project root directory dynamically.
    Works both in development and when frozen as .exe
    """
    if getattr(sys, 'frozen', False):
        # Running as compiled executable
        return os.path.dirname(sys.executable)
    else:
        # Running as script - find project root
        current = os.path.dirname(os.path.abspath(__file__))
        
        # Go up until we find project root indicators
        for _ in range(5):
            if os.path.exists(os.path.join(current, 'core')) or \
               os.path.exists(os.path.join(current, 'ui')) or \
               os.path.exists(os.path.join(current, 'main.py')):
                return current
            parent = os.path.dirname(current)
            if parent == current:
                break
            current = parent
        
        # Fallback
        return os.path.dirname(os.path.abspath(__file__))


class DocumentExporter:
    """Handles export of surveillance assignments to Word documents."""
    
    # Session time mappings
    SESSION_TIMES = {
        "S1": "08:30",
        "S2": "10:30",
        "S3": "12:30",
        "S4": "14:30"
    }
    
    # Reverse mapping for time to session code
    TIME_TO_SESSION = {v: k for k, v in SESSION_TIMES.items()}
    
    DURATION = "1h30"  # Standard exam duration
    MAX_ROWS_PER_PAGE = 20  # Adjust based on your page size
    
    def __init__(self, template_dir: str = "data/input/templates"):
        """
        Initialize the exporter with template directory.
        """
        try:
            # Use resource_path to get template directory dynamically
            self.template_dir = resource_path("data/input/templates")

            # Try different possible template filenames
            convocation_names = [
                "Exemple de convocation de suveillance.docx",
                "Exemple_de_convocation_de_surveillance.docx",
                "convocation_template.docx"
            ]

            affectation_names = [
                "Affectation des ensiegnants par jour --Surveillance .docx",
                "Affectation_des_enseignants_par_jour.docx",
                "affectation_template.docx"
            ]

            # === Find convocation template ===
            self.convocation_template = None
            for name in convocation_names:
                path = os.path.join(self.template_dir, name)
                if os.path.exists(path):
                    self.convocation_template = path
                    break

            # === Find affectation template ===
            self.affectation_template = None
            for name in affectation_names:
                path = os.path.join(self.template_dir, name)
                if os.path.exists(path):
                    self.affectation_template = path
                    break

            # === Sanity check ===
            if not self.convocation_template:
                raise FileNotFoundError(
                    f"Convocation template not found in {self.template_dir}. "
                    f"Tried: {convocation_names}"
                )

            if not self.affectation_template:
                raise FileNotFoundError(
                    f"Affectation template not found in {self.template_dir}. "
                    f"Tried: {affectation_names}"
                )

        except Exception as e:
            QMessageBox.critical(
                None,
                "Template Loading Error",
                f"Failed to load Word templates:\n{str(e)}"
            )
            sys.exit(1)

    
    def _get_academic_info(self) -> Dict[str, str]:
        """
        Calculate current academic year and semester based on current date.
        
        Returns:
            Dictionary with 'au' (academic year) and 'semester'
        """
        today = datetime.today()
        year = today.year
        month = today.month
        
        # Determine academic year
        if month >= 8:
            au = f"{year}-{year + 1}"
        else:
            au = f"{year - 1}-{year}"
        
        # Determine semester (2-7 = semester 2, else semester 1)
        semester = "2" if 2 <= month <= 7 else "1"
        
        return {"au": au, "semester": semester}
    
    def _format_date(self, date) -> str:
        """Format date for display."""
        if isinstance(date, str):
            return date
        return date.strftime("%d/%m/%Y")
    
    def _get_session_code(self, time_str: str) -> str:
        """Convert time string to session code (S1-S4)."""
        # Handle various time formats
        time_clean = str(time_str).strip().replace("h", ":").replace("H", ":")
        
        # Try direct mapping first
        if time_clean in self.TIME_TO_SESSION:
            return self.TIME_TO_SESSION[time_clean]
        
        # Try to extract hour
        try:
            hour = int(time_clean.split(":")[0])
            if hour == 8:
                return "S1"
            elif hour == 10:
                return "S2"
            elif hour == 12:
                return "S3"
            elif hour == 14:
                return "S4"
        except:
            pass
        
        return time_str  # Return as-is if can't convert
    
    def export_individual_convocations(
    self, 
    df_assignments: pd.DataFrame, 
    output_dir: str
) -> List[str]:
        """
        Generate individual convocation documents for each teacher from scratch,
        using only the template header/footer. Table and body are dynamically built.
        """
        os.makedirs(output_dir, exist_ok=True)
        generated_files = []
        academic_info = self._get_academic_info()

        grouped = df_assignments.groupby("Teacher")
        
        for teacher, df_teacher in grouped:
            try:
                # Create document from template (just header/footer)
                doc = Document(self.convocation_template)
                _fix_section_margins(doc)
                
                # --- Add intro paragraphs ---
                doc.add_paragraph(f"Notes à\nMr/Mme {teacher}")
                doc.add_paragraph(
                    "Cher(e) Collègue,\n"
                    "Vous êtes prié(e) d'assurer la surveillance et (ou) la responsabilité "
                    "des examens selon le calendrier ci-joint."
                )
                
                # --- Add table ---
                table = doc.add_table(rows=1, cols=3)
                set_table_borders(table)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Date"
                hdr_cells[1].text = "Heure"
                hdr_cells[2].text = "Durée"
                
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                
                # --- Sort and fill table ---
                df_sorted = df_teacher.sort_values(['Date', 'Time'])
                for _, row in df_sorted.iterrows():
                    new_row = table.add_row().cells
                    new_row[0].text = self._format_date(row['Date'])
                    new_row[1].text = str(row['Time'])
                    new_row[2].text = self.DURATION
                    for cell in new_row:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # --- Save document ---
                safe_name = teacher.replace(" ", "_").replace("/", "-")
                output_path = os.path.join(output_dir, f"Convocation_{safe_name}.docx")
                doc.save(output_path)
                generated_files.append(output_path)
            
            except Exception as e:
                import traceback
                print(f"Error generating convocation for {teacher}: {e}")
                print(traceback.format_exc())
                continue
        
        return generated_files
    
    def export_collective_affectations(
    self, 
    df_assignments: pd.DataFrame, 
    output_dir: str
) -> List[str]:
        """
        Generate collective affectation documents for each session from scratch.
        Only header/footer is kept from template; table is dynamically built.
        """
        os.makedirs(output_dir, exist_ok=True)
        generated_files = []
        academic_info = self._get_academic_info()

        grouped = df_assignments.groupby(['Date', 'Time'])
        
        for (date, time), df_session in grouped:
            try:
                # Create document from template (just header/footer)
                doc = Document(self.affectation_template)
                _fix_section_margins(doc)
                
                # --- Add header paragraph ---
                session_code = self._get_session_code(time)
                doc.add_paragraph(
                    f"AU : {academic_info['au']} – Semestre : {academic_info['semester']}\n"
                    f"Date : {self._format_date(date)} – Séance : {session_code}"
                )
                
                # --- Add table ---
                table = doc.add_table(rows=1, cols=3)
                set_table_borders(table)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Enseignant"
                hdr_cells[1].text = "Salle"
                hdr_cells[2].text = "Signature"
                
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                
                # --- Fill table with teachers ---
                df_sorted = df_session.sort_values('Teacher')
                for _, row in df_sorted.iterrows():
                    new_row = table.add_row().cells
                    new_row[0].text = str(row['Teacher'])
                    new_row[1].text = ""
                    new_row[2].text = ""
                    for i, cell in enumerate(new_row):
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if i > 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # --- Save document ---
                date_str = self._format_date(date).replace("/", "-")
                time_str = str(time).replace(":", "h")
                output_path = os.path.join(output_dir, f"Affectation_{date_str}_{time_str}.docx")
                doc.save(output_path)
                generated_files.append(output_path)
            
            except Exception as e:
                import traceback
                print(f"Error generating affectation for {date} {time}: {e}")
                print(traceback.format_exc())
                continue
        
        return generated_files

    
    def export_all(
        self, 
        df_assignments: pd.DataFrame, 
        output_dir: str,
        separate_dirs: bool = True
    ) -> Dict[str, List[str]]:
        """
        Export both individual convocations and collective affectations.
        
        Args:
            df_assignments: DataFrame with assignment data
            output_dir: Base output directory
            separate_dirs: If True, create separate subdirectories for each type
            
        Returns:
            Dictionary with keys 'convocations' and 'affectations' containing file lists
        """
        results = {}
        try:
            db = DBHandler()
            if hasattr(db, 'df') and isinstance(db.df, pd.DataFrame) and not db.df.empty:
                msg = QMessageBox()
                msg.setIcon(QMessageBox.Information)
                msg.setWindowTitle("Export Status")
                msg.setText("Using DataFrame from DBHandler (overriding provided data).")
                msg.exec_()
                df_assignments = db.df
            else:
                print("[ExportDocs] DBHandler has no valid DataFrame, using provided one.")
        except Exception as e:
            print(f"[ExportDocs] Failed to load DataFrame from DBHandler: {e}")
            
        if separate_dirs:
            convocation_dir = os.path.join(output_dir, "Convocations")
            affectation_dir = os.path.join(output_dir, "Affectations")
        else:
            convocation_dir = output_dir
            affectation_dir = output_dir
        
        # Generate individual convocations
        print("Generating individual convocations...")
        results['convocations'] = self.export_individual_convocations(
            df_assignments, 
            convocation_dir
        )
        print(f"Generated {len(results['convocations'])} convocation documents")
        
        # Generate collective affectations
        print("Generating collective affectations...")
        results['affectations'] = self.export_collective_affectations(
            df_assignments, 
            affectation_dir
        )
        print(f"Generated {len(results['affectations'])} affectation documents")
        
        return results


# Convenience function for easy integration
def export_surveillance_documents(
    df_assignments: pd.DataFrame,
    output_dir: str,
    template_dir: str = "data/templates"
) -> Dict[str, List[str]]:
    """
    Main export function - generates all surveillance documents.
    
    Args:
        df_assignments: DataFrame with columns ['Date', 'Time', 'Teacher']
        output_dir: Directory to save generated documents
        template_dir: Directory containing template files
        
    Returns:
        Dictionary with generated file paths
    """
    exporter = DocumentExporter(template_dir)
    return exporter.export_all(df_assignments, output_dir)


# Example usage
if __name__ == "__main__":
    # Example DataFrame structure
    sample_data = {
        'Date': ['13/05/2025', '13/05/2025', '14/05/2025'],
        'Time': ['08:30', '10:30', '08:30'],
        'Teacher': ['Prof. Ahmed', 'Prof. Ahmed', 'Prof. Fatma']
    }
    
    df = pd.DataFrame(sample_data)
    
    # Export documents
    results = export_surveillance_documents(
        df,
        output_dir="output/surveillance_docs",
        template_dir="data/templates"
    )
    
    print(f"\nExport complete!")
    print(f"Convocations: {len(results['convocations'])} files")
    print(f"Affectations: {len(results['affectations'])} files")

def _fix_section_margins(doc):
    for section in doc.sections:
        if not isinstance(section.left_margin, int):
            section.left_margin = int(section.left_margin)
        if not isinstance(section.right_margin, int):
            section.right_margin = int(section.right_margin)
        if not isinstance(section.top_margin, int):
            section.top_margin = int(section.top_margin)
        if not isinstance(section.bottom_margin, int):
            section.bottom_margin = int(section.bottom_margin)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # thickness
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)



